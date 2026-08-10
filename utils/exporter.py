import os
import re
import io
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor


def sanitize_filename(query: str) -> str:
    """Converts a user query into a clean filename."""
    clean = re.sub(r'[^\w\s-]', '', query).strip()
    return re.sub(r'[-\s]+', '_', clean)


def set_cell_background(cell, fill_hex: str):
    """Utility to set cell background color in python-docx."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def make_row_header(row):
    """Utility to make a table row repeat on every page in python-docx."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def parse_markdown_table_rows(lines: list[str]) -> tuple[list[list[str]], list[str]]:
    """Extracts table rows and returns (table_data, remaining_lines)."""
    table_data = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith("|") or not line.endswith("|"):
            break
        # Skip Markdown separator lines like |---|---|
        if re.match(r"^\|[\s:-]+(\+[\s:-]+)*\|$", line) or re.match(r"^\|(\s*:?-+:?\s*\|)+$", line):
            idx += 1
            continue

        row = [cell.strip() for cell in line.strip("|").split("|")]
        table_data.append(row)
        idx += 1

    return table_data, lines[idx:]


def process_docx_paragraph(paragraph, text: str):
    """Parses text containing markdown formatting (**bold**) into paragraph runs."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def generate_docx(report_text: str, search_query: str = "Research_Report") -> tuple[bytes, str]:
    """Generates a styled .docx file, returns bytes and filename for Streamlit downloading."""
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = report_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Check for Markdown Table Start
        if line.startswith("|") and line.endswith("|"):
            table_data, remaining_lines = parse_markdown_table_rows(lines[i:])
            lines = lines[:i] + remaining_lines

            if table_data:
                col_count = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=col_count)
                table.style = 'Table Grid'

                for row_idx, row in enumerate(table_data):
                    table_row = table.rows[row_idx]
                    
                    # Repeat header row across multi-page tables
                    if row_idx == 0:
                        make_row_header(table_row)

                    for col_idx, cell_value in enumerate(row):
                        if col_idx < col_count:
                            cell = table_row.cells[col_idx]
                            cell.text = ""  # Clear existing default run
                            p = cell.paragraphs[0]
                            
                            # Clean up inline tags and render formatted bold text inside cell
                            clean_val = cell_value.replace("<b>", "").replace("</b>", "")
                            process_docx_paragraph(p, clean_val)

                            # Format Header Row Appearance
                            if row_idx == 0:
                                set_cell_background(cell, "FF6B00")
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)

                doc.add_paragraph()  # Spacing after table
            continue

        # Headings (Strip '#' and formatting asterisks)
        if line.startswith("# "):
            doc.add_heading(re.sub(r"[\*#]", "", line).strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(re.sub(r"[\*#]", "", line).strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(re.sub(r"[\*#]", "", line).strip(), level=3)
        else:
            # Bullet vs Standard Paragraph
            is_bullet = line.startswith("- ") or line.startswith("* ") or line.startswith("• ")
            if is_bullet:
                raw_text = re.sub(r"^[-*•]\s*", "", line)
                p = doc.add_paragraph(style="List Bullet")
            else:
                raw_text = line
                p = doc.add_paragraph()

            # Render inline bold runs properly
            process_docx_paragraph(p, raw_text)

        i += 1

    file_name = f"{sanitize_filename(search_query)}.docx"

    # Save to buffer stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()

    return docx_bytes, file_name


def format_pdf_markdown(text: str) -> str:
    """Safely escapes HTML symbols and converts markdown bold (**) to ReportLab XML tags (<b>)."""
    # 1. Temporarily replace bold markers
    placeholders = []
    def replace_bold(match):
        placeholders.append(match.group(1))
        return f"__BOLD_PH_{len(placeholders)-1}__"

    text = re.sub(r"\*\*(.*?)\*\*", replace_bold, text)

    # 2. Escape XML special characters
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # 3. Restore bold tags safely for ReportLab Paragraph
    for idx, ph in enumerate(placeholders):
        # Escape any special characters inside the bold content itself
        clean_ph = ph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        text = text.replace(f"__BOLD_PH_{idx}__", f"<b>{clean_ph}</b>")

    return text


def generate_pdf(report_text: str, search_query: str = "Research_Report") -> tuple[bytes, str]:
    """Generates a formatted PDF file, returns bytes and filename for Streamlit downloading."""
    file_name = f"{sanitize_filename(search_query)}.pdf"
    printable_width = 504  # 612 letter width minus margins

    styles = getSampleStyleSheet()

    body_style = ParagraphStyle("ReportBody", parent=styles["Normal"], spaceAfter=8, fontSize=10, leading=14)
    h1_style = ParagraphStyle("ReportH1", parent=styles["Heading1"], spaceAfter=12, fontSize=18, leading=22, textColor=HexColor("#FF6B00"))
    h2_style = ParagraphStyle("ReportH2", parent=styles["Heading2"], spaceAfter=10, fontSize=14, leading=18, textColor=HexColor("#CC5500"))
    h3_style = ParagraphStyle("ReportH3", parent=styles["Heading3"], spaceAfter=8, fontSize=12, leading=15)
    bullet_style = ParagraphStyle("ReportBullet", parent=body_style, leftIndent=15, bulletIndent=5, spaceAfter=4)
    
    table_cell_style = ParagraphStyle("TableCell", parent=body_style, fontSize=9, leading=12, spaceAfter=0)
    table_header_style = ParagraphStyle("TableHeader", parent=table_cell_style, textColor=HexColor("#FFFFFF"), fontName="Helvetica-Bold")

    story = []
    lines = report_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            story.append(Spacer(1, 6))
            i += 1
            continue

        # Check for Markdown Table Start
        if line.startswith("|") and line.endswith("|"):
            table_data, remaining_lines = parse_markdown_table_rows(lines[i:])
            lines = lines[:i] + remaining_lines

            if table_data:
                col_count = max(len(row) for row in table_data)
                col_width = printable_width / col_count if col_count > 0 else printable_width
                col_widths = [col_width] * col_count

                formatted_table_data = []
                for row_idx, row in enumerate(table_data):
                    formatted_row = []
                    for cell_value in row:
                        clean_cell = format_pdf_markdown(cell_value)
                        style = table_header_style if row_idx == 0 else table_cell_style
                        formatted_row.append(Paragraph(clean_cell, style))
                    formatted_table_data.append(formatted_row)

                pdf_table = Table(
                    formatted_table_data, 
                    colWidths=col_widths, 
                    repeatRows=1
                )
                pdf_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#FF6B00")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 8),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#FFE0D1")),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#FFF5EE")]),
                        ]
                    )
                )
                story.append(pdf_table)
                story.append(Spacer(1, 10))
            continue

        # Escape raw HTML and process bold syntax safely
        formatted_line = format_pdf_markdown(line)

        if line.startswith("# "):
            clean_h1 = re.sub(r"^#\s*", "", line).replace("**", "")
            story.append(Paragraph(format_pdf_markdown(clean_h1), h1_style))
        elif line.startswith("## "):
            clean_h2 = re.sub(r"^##\s*", "", line).replace("**", "")
            story.append(Paragraph(format_pdf_markdown(clean_h2), h2_style))
        elif line.startswith("### "):
            clean_h3 = re.sub(r"^###\s*", "", line).replace("**", "")
            story.append(Paragraph(format_pdf_markdown(clean_h3), h3_style))
        elif line.startswith("- ") or line.startswith("* ") or line.startswith("• "):
            clean_bullet = re.sub(r"^[-*•]\s*", "", formatted_line)
            story.append(Paragraph(f"• {clean_bullet}", bullet_style))
        else:
            story.append(Paragraph(formatted_line, body_style))

        i += 1

    # Save into temp memory folder and return byte sequence
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = os.path.join(temp_dir, file_name)
        doc = SimpleDocTemplate(
            temp_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )
        doc.build(story)

        with open(temp_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes, file_name
